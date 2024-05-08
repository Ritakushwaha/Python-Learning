# -*- coding: utf-8 -*-
"""
Created on Thu Mar 25 17:45:57 2021

@author: Rita
"""

# import docx NOT python-docx
from docx import Document
from docx.table import _Cell
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import os

_date = str(datetime.today()).split()[0]

rows = 5
cols = 3
server = 'Server Name'
cpu = 'CPU Utilization %'
memory = 'Memory Utilization %'
status = 'Status'
status_value = 'OK'

s1 = '1'
s2 = '2'
s3 = '3'
s4 = '4'

mahti_cpu = 0
mahti_memory = 0
tilo_cpu = 0
tilo_memory = 0
marwa_cpu = 0
marwa_memory = 0

lst = [mahti_cpu, mahti_memory, tilo_cpu, tilo_memory, marwa_cpu, marwa_memory]

server_cpu = ['r', 'i', 't', 'a']
server_memory = ['g', 'i', 'r', 'u']

# create an instance of a word document
doc = Document()

# add a heading of level 0 (largest heading)
doc.add_heading('Capacity report for the week', 0)

# add a paragraph and store
# the object in a variable
doc_para = doc.add_paragraph()

# add a run i.e, style like
# bold, italic, underline, etc.
doc_para.add_run('CPU STATUS').bold = True
doc_para.add_run('\nBelow is the current CPU utilization.')

table1 = doc.add_table(rows, cols)
table1.style = 'Table Grid'

# populate header row --------
heading_cells = table1.rows[0].cells
heading_cells[0].text = server
heading_cells[1].text = cpu
heading_cells[2].text = status

table1.cell(1, 0).text = s1
table1.cell(2, 0).text = s2
table1.cell(3, 0).text = s3
table1.cell(4, 0).text = s4

table1.cell(1, 1).text = server_cpu[0]
table1.cell(2, 1).text = server_cpu[1]
table1.cell(3, 1).text = server_cpu[2]
table1.cell(4, 1).text = server_cpu[3]

for i in range(0, rows):
    table1.cell(i, 2).text = status_value

doc_para1 = doc.add_paragraph()

doc_para1.add_run('\n\nMEMORY STATUS').bold = True
doc_para1.add_run('\nMemory utilization trend is as follows.')

table2 = doc.add_table(rows, cols)
table2.style = 'Table Grid'

# populate header row --------
heading_cells = table2.rows[0].cells
heading_cells[0].text = server
heading_cells[1].text = memory
heading_cells[2].text = status

table2.cell(1, 0).text = s1
table2.cell(2, 0).text = s2
table2.cell(3, 0).text = s3
table2.cell(4, 0).text = s4

table2.cell(1, 1).text = server_memory[0]
table2.cell(2, 1).text = server_memory[1]
table2.cell(3, 1).text = server_memory[2]
table2.cell(4, 1).text = server_memory[3]

for i in range(0, rows):
    table2.cell(i, 2).text = status_value

doc_para2 = doc.add_paragraph()

doc_para2.add_run('\n\n{s1} CPU: OK').bold = True
doc_para2.add_run(f'\nThe CPU utilization for {s1} is around {lst[0]}%.')

doc_para2.add_run('\n\n{s1} MEMORY: OK').bold = True
doc_para2.add_run(f'\nThe memory utilization for {s1} is currently around {lst[1]}%.')

doc_para2.add_run('\n\n{s2} FIN CPU: OK').bold = True
doc_para2.add_run(f'\nThe CPU utilization for {s2} is around {lst[2]}%.')

doc_para2.add_run('\n\n{s2} FIN MEMORY: OK').bold = True
doc_para2.add_run(f'\nThe memory utilization for {s2} is currently around {lst[3]}%.')

doc_para2.add_run('\n\n{s4} CPU: OK').bold = True
doc_para2.add_run(f'\nThe CPU utilization for {s4} is around {lst[4]}%. ')

doc_para2.add_run('\n\n{s4} MEMORY: OK').bold = True
doc_para2.add_run(f'\nThe memory utilization for {s4} is currently around {lst[5]}%')

doc_para2.add_run('\n\n\nHello: ').bold = True
doc_para2.add_run('\nHi').bold = True
doc_para2.add_run('\nSome text.')

doc_para2.add_run('\n\nYou').bold = True
doc_para2.add_run('\nSome text.')

doc_para2.add_run('\n\nHey').bold = True
doc_para2.add_run('\nListen')

# now save the document to a location
doc.save(f'test_word.docx')
